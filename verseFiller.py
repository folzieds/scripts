import docx
import pandas as pd
import re


df_verse = pd.read_csv("verse.csv")

def get_text(filename: str) -> list:
    doc = docx.Document(filename)

    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    return full_text

def fill_verse(full_text: list) -> list:
    filled_text = []
    for text in full_text:
        if(re.search('[1-3]*\s[A-Za-z]*\s[0-9]*:\s[1-9]*\s-\s[0-9]*', text) != None):
            temp = add_range_verse_text(re.search('[1-3]*\s[A-Za-z]*\s[0-9]*:\s[0-9]*\s-\s[0-9]*', text).group())
            filled_text.append(temp)
        elif (re.search('[A-Za-z]*\s[0-9]*:\s[0-9]*\s-\s[0-9]*', text)) != None:
            temp = add_range_verse_text(re.search('[A-Za-z]*\s[0-9]*:\s[0-9]*\s-\s[0-9]*', text).group())
            filled_text.append(temp)
        elif(re.search('[1-3]*\s[A-Za-z]*\s[0-9]*:\s[0-9]*', text) != None):
            temp = add_verse_text(re.search('[1-3]*\s[A-Za-z]*\s[0-9]*:\s[0-9]*', text).group())
            filled_text.append(temp)
        elif (re.search('[A-Za-z]*\s[0-9]*:\s[0-9]*', text)) != None:
            temp = add_verse_text(re.search('[A-Za-z]*\s[0-9]*:\s[0-9]*', text).group())
            filled_text.append(temp)
        else:
            filled_text.append(text)
        
    return filled_text

def add_verse_text(verse: str) -> str:
    # split the verse into book chapter and verse
    book = ""
    if(re.search("[0-9]*\s[a-zA-z]*",verse).group().strip() != ''):
        book = re.search("[0-9]*\s[a-zA-z]*",verse).group().strip()
    else:
        book = re.search("[a-zA-z]*",verse).group().strip()
    chapter = int(re.split("[a-zA-z]",verse.split(":",maxsplit = 2)[0])[-1].strip())
    ver = int(verse.split(":",maxsplit = 2)[-1].strip())
    # get the text from the dataframe
    verse_df = df_verse.loc[(df_verse['book'] == book) & (df_verse['chapter'] == chapter) & (df_verse['verse'] == ver)]
    verse_text = verse_df.iat[0,3]

    return f"{book} {chapter}: {ver} - {verse_text}"

def add_range_verse_text(verse: str) -> str:
    # split the verse into book chapter and verse
    book = ""
    if(re.search("[0-9]*\s[a-zA-z]*",verse).group().strip() != ''):
        book = re.search("[0-9]*\s[a-zA-z]*",verse).group().strip()
    else:
        book = re.search("[a-zA-z]*",verse).group().strip()
    chapter = int(re.split("[a-zA-z]",verse.split(":",maxsplit = 2)[0])[-1].strip())
    ver_start = int(verse.split(":",maxsplit = 2)[1].strip().split("-")[0].strip())
    ver_end = int(verse.split(":",maxsplit = 2)[1].strip().split("-")[1].strip())

    verse_list = []
    # get the text from the dataframe
    for ver in range(ver_start,ver_end+1):
        verse_df = df_verse.loc[(df_verse['book'] == book) & (df_verse['chapter'] == chapter) & (df_verse['verse'] == ver)]
        verse_text = verse_df.iat[0,3]
        verse_compile = f"{book} {chapter}: {ver} - {verse_text}"
        verse_list.append(verse_compile)

    return "\n".join(verse_list)

def export_text_to_docx(text_list :list):
    doc = docx.Document()
    #parse list to a single string
    for text in text_list:
        doc.add_paragraph(text)
    
    doc.save("filled.docx")

x = get_text("test.docx")
y = fill_verse(x)
export_text_to_docx(y)
